#!/usr/bin/env python3
"""Extrae párrafos y tablas básicas de un .docx del PEI hacia texto/markdown auxiliar.

Uso:
  python tools/extract_pei_docx.py /ruta/al/PEI.docx > tmp/pei_dump.md

No sustituye la curaduría humana para armar `config/plan_*.yaml` oficial.
"""
from __future__ import annotations

import pathlib
import sys


def main() -> None:
    if len(sys.argv) < 2:
        print("Uso: extract_pei_docx.py <archivo.docx>", file=sys.stderr)
        sys.exit(1)
    path = pathlib.Path(sys.argv[1])
    if not path.is_file():
        print("No existe:", path, file=sys.stderr)
        sys.exit(1)
    try:
        import docx  # python-docx
    except ImportError:
        print("Instale python-docx: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    document = docx.Document(str(path))
    print(f"# Extracción automática: {path.name}\n")
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            print(t + "\n")
    for ti, table in enumerate(document.tables):
        print(f"\n## Tabla {ti + 1}\n")
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            print("| " + " | ".join(cells) + " |")


if __name__ == "__main__":
    main()

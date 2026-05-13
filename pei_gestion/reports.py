"""Exportación Excel y Word con bloque de metadatos."""
from __future__ import annotations

import datetime as dt
import pathlib
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt


def _meta_block(plan_name: str, source: str) -> list[tuple[str, str]]:
    now = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    return [
        ("Plan", plan_name),
        ("Fuente de datos", source),
        ("Generado", now),
        ("Sistema", "PEI UCCuyo — Streamlit"),
    ]


def write_excel_report(
    path: str | pathlib.Path,
    *,
    long_df: pd.DataFrame,
    by_unit_year: pd.DataFrame,
    by_og: pd.DataFrame,
    completeness: pd.DataFrame,
    plan_name: str,
    source: str,
    validated: pd.DataFrame | None = None,
    layer_b: pd.DataFrame | None = None,
    meta_kpis: dict[str, Any] | None = None,
) -> None:
    path = pathlib.Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    meta = pd.DataFrame(_meta_block(plan_name, source), columns=["campo", "valor"])
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        meta.to_excel(w, sheet_name="metadata", index=False)
        long_df.to_excel(w, sheet_name="actividades_long", index=False)
        by_unit_year.to_excel(w, sheet_name="resumen_unidad_anio", index=False)
        by_og.to_excel(w, sheet_name="resumen_og", index=False)
        completeness.to_excel(w, sheet_name="completitud_flags", index=False)
        if validated is not None and not validated.empty:
            validated.to_excel(w, sheet_name="validacion_oe_unidad", index=False)
        if layer_b is not None and not layer_b.empty:
            layer_b.to_excel(w, sheet_name="capaB_oe_matriz", index=False)
        if meta_kpis:
            pd.DataFrame([meta_kpis]).to_excel(w, sheet_name="capaC_meta_kpis", index=False)
        if not long_df.empty and "email" in long_df.columns and "unidad" in long_df.columns:
            lf = long_df.copy()
            lf["email"] = lf["email"].fillna("sin_email").astype(str)
            pivot = (
                lf.pivot_table(index="email", columns="unidad", values="respuesta_id", aggfunc="count", fill_value=0)
                .astype(int)
            )
            if pivot.shape[1] > 80:
                pivot = pivot.iloc[:, :80]
            pivot.to_excel(w, sheet_name="pivot_email_unidad")


def write_word_summary(
    path: str | pathlib.Path,
    *,
    long_df: pd.DataFrame,
    plan_name: str,
    source: str,
    max_table_rows: int = 200,
) -> None:
    path = pathlib.Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    h = doc.add_heading("Informe de actividades — PEI", level=0)
    h.runs[0].font.size = Pt(16)
    doc.add_paragraph(f"Plan: {plan_name}")
    doc.add_paragraph(f"Fuente: {source}")
    doc.add_paragraph(f"Fecha de generación: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Resumen cuantitativo", level=1)
    doc.add_paragraph(f"Total de registros de actividad (formato largo): {len(long_df)}")
    if not long_df.empty:
        doc.add_paragraph(f"Unidades distintas: {long_df['unidad'].nunique()}")
        doc.add_paragraph(f"Responsables (email) distintos: {long_df['email'].nunique()}")

    doc.add_heading("Muestra de actividades", level=1)
    doc.add_paragraph(
        "La tabla siguiente lista hasta {0} filas con unidad, año, OG, texto de OE y actividad.".format(
            max_table_rows
        )
    )
    sample = long_df.head(max_table_rows)
    table = doc.add_table(rows=1, cols=8)
    hdr = table.rows[0].cells
    cols = ["Unidad", "Año", "OG", "OE", "Actividad", "Resultado", "Origen", "Acción/Indicador"]
    for i, name in enumerate(cols):
        hdr[i].text = name
    for _, r in sample.iterrows():
        row = table.add_row().cells
        row[0].text = str(r.get("unidad", ""))
        row[1].text = str(r.get("anio", ""))
        row[2].text = str(r.get("og", ""))
        row[3].text = str(r.get("oe", ""))[:500]
        row[4].text = str(r.get("actividad", ""))[:500]
        row[5].text = str(r.get("resultado", ""))[:500]
        row[6].text = str(r.get("origen", ""))
        row[7].text = (str(r.get("accion", "")) + " / " + str(r.get("indicador", "")))[:500]

    doc.add_heading("Indicadores formales del PEI (documento)", level=1)
    doc.add_paragraph(
        "Las métricas de indicadores del libro del PEI solo aplican cuando existen valores en fuentes institucionales "
        "o registros con acción/indicador explícitos (p. ej. carga directa en la app). "
        "El instrumento Google Forms vigente certifica hasta el nivel de objetivo específico (OE)."
    )

    doc.save(path)
